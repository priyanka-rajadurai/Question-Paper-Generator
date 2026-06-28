from docx import Document
from flask import Flask, render_template, request
import pandas as pd
import os
import io
import tempfile
from PIL import Image
from docx.shared import Inches

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

@app.route("/", methods=["GET", "POST"])
def home():

    message = ""

    if request.method == "POST":
        print(request.form)
        print("Description =", request.form.get("part_description[]"))

        file = request.files.get("excel_file")

        if file and file.filename:

            filepath = os.path.join(
                app.config["UPLOAD_FOLDER"],
                file.filename
            )

            file.save(filepath)
            from openpyxl import load_workbook

            # Load the uploaded Excel file
            wb = load_workbook(filepath)
            ws = wb.active

            # Store all embedded images
            image_dict = {}

            for img in ws._images:
                row = img.anchor._from.row + 1   # Excel row number
                image_dict[row] = img

            print("Images Found:", image_dict.keys())
            print("\n===== QUESTIONS THAT HAVE IMAGES =====")

            for row_no in sorted(image_dict.keys()):
                try:
                    print(
                        "Excel Row:", row_no,
                        "| S.No:", ws.cell(row=row_no, column=1).value,
                        "| Question:",
                        str(ws.cell(row=row_no, column=2).value)[:80]
                    )
                except Exception as e:
                    print(row_no, e)

            print("=====================================\n")
            df = pd.read_excel(
                filepath,
                engine="openpyxl"
            )
            # Store original Excel row number
            df["Excel_Row"] = df.index + 2
            print(df[["Question", "Excel_Row"]].head())
            print(df.columns)

            part_name = request.form.getlist("part_name[]")
            part_marks = request.form.getlist("part_marks[]")
            part_co = request.form.getlist("part_co[]")
            part_questions = request.form.getlist("part_questions[]")
            part_description = request.form.getlist("part_description[]")
            optional = request.form.getlist("optional[]")
            print(part_name)
            print(part_marks)
            print(part_co)
            print(part_questions)
            if part_name:


                doc = Document()
                from docx.shared import Cm

                section = doc.sections[0]

                section.page_width = Cm(21.0)      
                section.page_height = Cm(29.7)     

                section.left_margin = Cm(2.0)
                section.right_margin = Cm(2.0)
                section.top_margin = Cm(1.27)
                section.bottom_margin = Cm(1.27)

                from docx.shared import Inches

                section = doc.sections[0]
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                    
                table.cell(0,0).text = "Question Paper Code"
                table.cell(0,1).text = request.form.get("qp_code")

                table.cell(0,2).text = "Register No"
                table.cell(0,3).text = ""
                doc.add_paragraph()

                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                p = doc.add_paragraph()

                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = p.add_run(request.form.get("institution"))

                run.bold = True
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                p = doc.add_paragraph()

                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = p.add_run(request.form.get("exam_name"))

                run.bold = True
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                p = doc.add_paragraph()

                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = p.add_run(request.form.get("semester"))

                run.bold = True

                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                p = doc.add_paragraph()

                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = p.add_run(request.form.get("exam_no"))

                run.bold = True

                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                p = doc.add_paragraph()

                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = p.add_run(request.form.get("department"))

                run.bold = True

                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = p.add_run(
                    f"{request.form.get('subject_code')} - {request.form.get('subject_name')}"
                )
                run.bold = True

                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                p.add_run(
                    f"Regulation: {request.form.get('regulation')}"
                ).bold = True
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                from docx.shared import Cm
                info_table = doc.add_table(rows=1, cols=2)
                info_table.autofit = False

                info_table.columns[0].width = Cm(10)
                info_table.columns[1].width = Cm(7)

                left = info_table.cell(0, 0).paragraphs[0]
                left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                left.add_run(f"Duration : {request.form.get('duration')}")

                right = info_table.cell(0, 1).paragraphs[0]
                right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                right.add_run(f"Max Marks : {request.form.get('max_marks')}")
                right = info_table.cell(0, 1).paragraphs[0]
            
                doc.add_paragraph()
                co_marks = {
                    "CO1":0,
                    "CO2":0,
                    "CO3":0,
                    "CO4":0,
                    "CO5":0
                }

                for i in range(len(part_name)):
                    marks = int(part_marks[i])
                    questions = int(part_questions[i])

                    total = marks * questions

                    co_marks[part_co[i]] += total
                q_no = 1
                for i in range(len(part_name)):

                    filtered = df[
                        (df["Part"] == part_name[i]) &
                        (df["Marks"] == int(part_marks[i])) &
                        (df["CO"] == part_co[i])
                    ]
                    is_optional = optional[i]

                    required = int(part_questions[i])

                    if is_optional == "YES":
                        total_required = required * 2
                    else:
                        total_required = required

                    if len(filtered) < total_required:
                        print(f"Only {len(filtered)} questions available for Part {part_name[i]}")
                        continue
                    print("Part:", part_name[i])
                    print("Filtered Questions:", len(filtered))
                    print("Required:", required)
                    print("Total Required:", total_required)

                    selected_questions = filtered.sample(n=total_required)
                    print("ID after sample:", id(selected_questions))
                    print(selected_questions)
                    print("After sample")
                    print(selected_questions.shape)
                    print(selected_questions[["Question", "Excel_Row"]])
                    print(type(filtered))
                    print(filtered.shape)
                    print("\n===== Selected Questions with Excel Rows =====")
                    print(selected_questions[["S.No", "Question", "Excel_Row"]])
                    print("=============================================")

                    print("------ Selected Questions ------")
                    print("ID before loop:", id(selected_questions))
                    print(selected_questions)
                    for idx, row in selected_questions.iterrows():
                        print(
                            "DF Index:", idx,
                            "| S.No:", row["S.No"],
                            "| Question:", row["Question"][:40]
                        )

                    print("-------------------------------")
                    print("Questions Selected:", len(selected_questions))

                    # Part Heading
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    run = p.add_run(f"PART {part_name[i]}")
                    run.bold = True

                    # Description
                    if part_description[i]:

                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p.add_run(part_description[i])

                    # Question Table

                    q_table = doc.add_table(rows=0, cols=5)
                    q_table.autofit = False
                    from docx.shared import Cm

                    q_table.columns[0].width = Cm(1.3)
                    q_table.columns[1].width = Cm(12.5)
                    q_table.columns[2].width = Cm(1.2)
                    q_table.columns[3].width = Cm(1.2)
                    q_table.columns[4].width = Cm(1.2)

                    if is_optional == "NO":
                        for idx, row in selected_questions.iterrows():
                            print("DataFrame Index:", idx)
                            cells = q_table.add_row().cells
                            cells[0].width = Cm(1.2)
                            cells[1].width = Cm(12.5)
                            cells[2].width = Cm(1.2)
                            cells[3].width = Cm(1.2)
                            cells[4].width = Cm(1.2)

                            cells[0].text = str(q_no)
                            cells[1].text = str(row["Question"])
                            excel_row = row["Excel_Row"]

                            if excel_row in image_dict:
                                print("Image found:", excel_row)

                                img = image_dict.get(row["Excel_Row"])

                                if img:
                                    image_bytes = img._data()

                                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                                    temp_file.write(image_bytes)
                                    temp_file.close()

                                    doc.add_picture(temp_file.name, width=Inches(3.5))

                                # Get image bytes
                                image_bytes = img._data()

                                # Save temporarily
                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                                temp_file.write(image_bytes)
                                temp_file.close()

                                # Insert image below the table
                                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                                p = doc.add_paragraph()
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                                run = p.add_run()
                                run.add_picture(temp_file.name, width=Inches(3.5))
                            cells[2].text = str(row["Marks"])
                            cells[3].text = str(row["CO"])
                            cells[4].text = str(row["K"])

                            q_no += 1

                    else:

                        print(">>> OPTIONAL BLOCK <<<")

                        questions = list(selected_questions.iterrows())

                        for idx, row in selected_questions.iterrows():
                            print("DataFrame Index:", idx)
                        
                        for j in range(0, len(questions), 2):
                            idx1, q1 = questions[j]
                            idx2, q2 = questions[j + 1]

                            excel_row1 = q1["Excel_Row"]
                            excel_row2 = q2["Excel_Row"]

                            print("Question A Excel Row:", excel_row1)
                            print("Question B Excel Row:", excel_row2)
                            if excel_row1 in image_dict:
                                print("Image found for A")

                            if excel_row2 in image_dict:
                                print("Image found for B")
                            # -------- Question (a) --------
                            q_table = doc.add_table(rows=1, cols=5)
                            q_table.autofit = False

                            cells = q_table.rows[0].cells

                            # Set widths for EACH CELL
                            cells[0].width = Cm(1.2)
                            cells[1].width = Cm(12.5)
                            cells[2].width = Cm(1.2)
                            cells[3].width = Cm(1.2)
                            cells[4].width = Cm(1.2)

                            cells[0].text = f"{q_no}(a)"
                            cells[1].text = str(q1["Question"])
                            if q1["Excel_Row"] in image_dict:
                                print("Image found for A")

                            img = image_dict.get(q1["Excel_Row"])

                            if img:
                                image_bytes = img._data()

                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                                temp_file.write(image_bytes)
                                temp_file.close()

                                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                                p = doc.add_paragraph()
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                                run = p.add_run()
                                run.add_picture(temp_file.name, width=Inches(3.5))
                            cells[2].text = str(q1["Marks"])
                            cells[3].text = str(q1["CO"])
                            cells[4].text = str(q1["K"])

                            # -------- OR --------
                            p = doc.add_paragraph()
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            p.add_run("OR").bold = True

                            # -------- Question (b) --------
                            q_table = doc.add_table(rows=1, cols=5)
                            q_table.autofit = False

                            cells = q_table.rows[0].cells

                            # Set widths for EACH CELL
                            cells[0].width = Cm(1.2)
                            cells[1].width = Cm(12.5)
                            cells[2].width = Cm(1.2)
                            cells[3].width = Cm(1.2)
                            cells[4].width = Cm(1.2)

                            cells[0].text = f"{q_no}(b)"
                            cells[1].text = str(q2["Question"])
                            if q2["Excel_Row"] in image_dict:
                                print("Image found for B")

                                img = image_dict.get(q2["Excel_Row"])

                                if img:
                                    image_bytes = img._data()

                                    # Save temporarily
                                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                                    temp_file.write(image_bytes)
                                    temp_file.close()

                                    # Insert image below the table
                                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                                    p = doc.add_paragraph()
                                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                                    run = p.add_run()
                                    run.add_picture(temp_file.name, width=Inches(3.5))
                            cells[2].text = str(q2["Marks"])
                            cells[3].text = str(q2["CO"])
                            cells[4].text = str(q2["K"])

                            q_no += 1
                    doc.add_paragraph()
                
               
                doc.add_paragraph()
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                run = p.add_run("Distribution of CO's (Percentage wise)")
                run.bold = True

                table = doc.add_table(rows=3, cols=6)
                table.style = "Table Grid"

                # First Row
                table.cell(0,0).text = "Evaluation"
                table.cell(0,1).text = "CO1"
                table.cell(0,2).text = "CO2"
                table.cell(0,3).text = "CO3"
                table.cell(0,4).text = "CO4"
                table.cell(0,5).text = "CO5"

                # Second Row
                table.cell(1,0).text = "Marks"

                # Third Row
                table.cell(2,0).text = "%"
                for r in [1,2]:
                    for c in range(1,6):
                        table.cell(r,c).text = "-"
                max_marks = int(request.form.get("max_marks"))

                co_list = ["CO1", "CO2", "CO3", "CO4", "CO5"]

                for col, co in enumerate(co_list, start=1):
                    if co_marks[co] != 0:
                        table.cell(1, col).text = str(co_marks[co])

                        percent = (co_marks[co] / max_marks) * 100
                        table.cell(2, col).text = f"{percent:.0f}"
                    else:
                        table.cell(1, col).text = "-"
                        table.cell(2, col).text = "-"
                doc.add_paragraph()

                p = doc.add_paragraph()

                p.add_run(
                    "Knowledge level: "
                    "K1 - Remember, "
                    "K2 - Understand, "
                    "K3 - Apply, "
                    "K4 - Analyse, "
                    "K5 - Evaluate, "
                    "K6 - Create"
                )
                print("Word file generated successfully")
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                return send_file(
                    buffer,
                    as_attachment=True,
                    download_name="Question_Paper.docx",
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                    
            message = f"{len(df)} Questions Loaded Successfully"

    return render_template(
        "index.html",
        message=message
    )

if __name__ == "__main__":
    app.run(debug=True)
